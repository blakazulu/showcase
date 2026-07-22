#!/usr/bin/env python
"""Convert a CV markdown file (see references/cv-contract.md) to an ATS-safe DOCX.

Usage:
    python md_to_docx.py cv/liraz-en.md            # -> cv/liraz-en.docx
    python md_to_docx.py cv/liraz-he.md out.docx   # explicit output

One column, standard headings, contact block as plain-text paragraphs at the top
(never a Word header/footer), no tables/text boxes/graphics. After writing, it
prints the DOCX's extracted plain text so you can eyeball reading order and confirm
the contact line survived (the guide's "export to TXT and inspect" ATS self-test).
"""
import sys
import os

# Windows consoles default to cp1252; force UTF-8 so the Hebrew self-test dump
# prints instead of crashing.
try:
    sys.stdout.reconfigure(encoding="utf-8")
    sys.stderr.reconfigure(encoding="utf-8")
except Exception:
    pass

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit(
        "Missing dependency 'python-docx'. Install it with:\n    pip install python-docx"
    )

from cv_parse import parse_cv, AI_TELL_CHARS  # local module next to this script


NEAR_BLACK = RGBColor(0x1A, 0x1A, 0x1A)
GREY = RGBColor(0x44, 0x44, 0x44)


def _set_rtl(paragraph):
    """Mark a paragraph (and its runs) right-to-left for Hebrew."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)
    for run in paragraph.runs:
        rPr = run._r.get_or_add_rPr()
        rtl = OxmlElement("w:rtl")
        rPr.append(rtl)


def _set_narrow_margins(document):
    for section in document.sections:
        section.top_margin = Pt(36)     # 0.5"
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(50)    # ~0.7"
        section.right_margin = Pt(50)


def _para(document, rtl, space_before=0, space_after=2):
    p = document.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.12
    if rtl:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return p


def _run(p, text, bold=False, size=10.5, color=NEAR_BLACK):
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = "Calibri"
    r.font.color.rgb = color
    return r


def build(cv, out_path):
    rtl = cv["dir"] == "rtl"
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)
    _set_narrow_margins(doc)

    fm = cv["front"]

    # --- header block (plain-text paragraphs, top of body) ---
    p = _para(doc, rtl, space_after=1)
    _run(p, fm.get("name", ""), bold=True, size=18)
    if rtl:
        _set_rtl(p)

    if fm.get("title"):
        p = _para(doc, rtl, space_after=3)
        _run(p, fm["title"], size=12, color=GREY)
        if rtl:
            _set_rtl(p)

    contact_bits = [b for b in (fm.get("location"), fm.get("phone"), fm.get("email")) if b]
    if contact_bits:
        p = _para(doc, rtl, space_after=1)
        _run(p, "  |  ".join(contact_bits), size=10, color=GREY)
        if rtl:
            _set_rtl(p)

    if cv["links"]:
        p = _para(doc, rtl, space_after=6)
        _run(p, "  |  ".join(f"{lbl}: {url}" for lbl, url in cv["links"]), size=9.5, color=GREY)
        # links stay LTR even in a Hebrew CV (URLs are LTR)

    # --- sections ---
    for section in cv["sections"]:
        p = _para(doc, rtl, space_before=8, space_after=2)
        _run(p, section["heading"].upper(), bold=True, size=11)
        if rtl:
            _set_rtl(p)
        # thin rule under heading via bottom border
        _heading_border(p)

        for block in section["blocks"]:
            if block["type"] == "entry":
                ep = _para(doc, rtl, space_before=4, space_after=1)
                fields = block["fields"]
                _run(ep, fields[0], bold=True, size=10.5)
                if len(fields) > 1:
                    _run(ep, "  |  " + "  |  ".join(fields[1:]), size=10, color=GREY)
                if rtl:
                    _set_rtl(ep)
            elif block["type"] == "bullet":
                bp = doc.add_paragraph(style="List Bullet")
                bp.paragraph_format.space_after = Pt(1)
                bp.paragraph_format.line_spacing = 1.12
                _run(bp, block["text"], size=10.5)
                if rtl:
                    bp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    _set_rtl(bp)
            elif block["type"] == "para":
                pp = _para(doc, rtl, space_after=2)
                _run(pp, block["text"], size=10.5)
                if rtl:
                    _set_rtl(pp)

    doc.save(out_path)
    return doc


def _heading_border(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "999999")
    pbdr.append(bottom)
    pPr.append(pbdr)


def self_test_dump(out_path, fm):
    """Re-read the DOCX and print extracted text; flag missing contact info."""
    doc = Document(out_path)
    lines = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(lines)
    print("\n----- ATS self-test: extracted DOCX text -----")
    print(text)
    print("----- end extracted text -----")
    missing = []
    for key in ("phone", "email"):
        val = fm.get(key, "")
        if val and val not in text:
            missing.append(f"{key} ({val})")
    if missing:
        print("WARNING: contact info NOT found in extracted text: " + ", ".join(missing))
    else:
        print("OK: contact info present in extracted text.")

    tells = sorted({ch for ch in text if ch in AI_TELL_CHARS})
    if tells:
        print("WARNING: AI-tell characters survived: " + " ".join(f"U+{ord(c):04X}" for c in tells))
    else:
        print("OK: no AI-tell punctuation (plain hyphens/quotes only).")


def main():
    if len(sys.argv) < 2:
        sys.exit("Usage: python md_to_docx.py <cv.md> [out.docx]")
    src = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else os.path.splitext(src)[0] + ".docx"
    with open(src, encoding="utf-8") as f:
        cv = parse_cv(f.read())
    build(cv, out)
    print(f"Wrote {out}")
    self_test_dump(out, cv["front"])


if __name__ == "__main__":
    main()
